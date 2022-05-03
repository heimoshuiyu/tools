import openpyxl
import docx
import os
import warnings


def main():
    doc_filename = input('输入报表附注模板：').strip()
    doc_filename = clean_filename(doc_filename)
    if not doc_filename.endswith('.docx'):
        print('【警告】：这不是 .docx 结尾的文件')

    wb_filename_list = []
    print('输入国资委和年度决算报表，可直接拖入文件夹，多个Excel将自动合并')
    while True:
        wb_filename = input('输入工作簿，留空以停止输入：').strip()
        wb_filename = clean_filename(wb_filename)
        if not wb_filename:
            break
        if os.path.isdir(wb_filename):
            for i in os.listdir(wb_filename):
                if i.endswith('.xlsx') and not (i.startswith('.') or i.startswith('~')):
                    wb_filename_list.append(os.path.join(wb_filename, i))
        else:
            wb_filename_list.append(wb_filename)

    output_filename = input('输出附注文件名：').strip()
    output_filename = clean_filename(output_filename)
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'

    with warnings.catch_warnings(record=True) as _:
        warnings.simplefilter("always")
        print('读取工作簿文件：', wb_filename_list[0])
        wb = openpyxl.load_workbook(wb_filename_list[0], data_only=True)
        # merge sheets
        for wb_filename in wb_filename_list[1:]:
            print('合并工作簿：', wb_filename)
            new_wb = openpyxl.load_workbook(wb_filename, data_only=True)
            append_wb(wb, new_wb)

    print('加载文件', doc_filename)
    doc = docx.Document(doc_filename)

    # print('加载文件', wb_filename)
    # wb = openpyxl.load_workbook(wb_filename, data_only=True)

    print('取消所有合并单元格...')
    unmerge_all_cells(wb)

    print('删除所有空行...')
    remove_all_empty_rows(wb)

    print('执行替换...')
    replace_cells_in_table(doc, wb)

    print('保存文件...')
    # save file
    # wb.save('/tmp/out.xlsx')
    doc.save(output_filename)

    input('按任意键退出...')


def append_wb(wb, new_wb):
    total_sheet_nums = len(new_wb.sheetnames)
    for count, sheet_name in enumerate(new_wb.sheetnames):
        print('Processing sheet', count, '/', total_sheet_nums, end='\r')
        new_sheet_name = sheet_name
        if sheet_name in wb.sheetnames:
            new_sheet_name = sheet_name + '-1'
        wb.create_sheet(new_sheet_name)
        for row in new_wb[sheet_name].rows:
            for cell in row:
                wb[new_sheet_name][cell.coordinate].value = cell.value
    print()


def clean_filename(s):
    s = s.strip()
    if s.startswith('"'):
        s = s[1:]
    if s.endswith('"'):
        s = s[:-1]
    if s.startswith("'"):
        s = s[1:]
    if s.endswith("'"):
        s = s[:-1]
    return s.strip()


def is_number(i):
    try:
        float(i)
        return True
    except ValueError:
        return False

# add thousands separator to numbers


def format_cell(v):
    """
    Format the cell to be a string
    """
    if v is None:
        return ''
    if isinstance(v, str):
        return v
    elif is_number(v):
        # must 2 digit after decimal point
        return '{:,.02f}'.format(v)
    else:
        return str(v)


def unmerge_all_cells(wb):
    """
    Unmerge all cells in all sheets
    """
    for sheet in wb:
        # unmerge all cells
        for cg in list(sheet.merged_cells):
            sheet.unmerge_cells(str(cg))


def remove_all_empty_rows(wb):
    """
    Remove all empty rows in all sheets
    """
    for sheet in wb:
        # remove empty rows
        for row in sheet.iter_rows(min_row=sheet.min_row, max_row=sheet.max_row, min_col=sheet.min_column, max_col=sheet.max_column):
            if all(cell.value is None for cell in row):
                sheet.delete_rows(row[0].row)


def replace_cells_in_table(doc, wb):
    total = len(doc.tables)
    for count, table in enumerate(doc.tables):
        print('Processing table', count+1, '/', total, end='\r')
        for row in table.rows:
            for cell in row.cells:
                celltext = cell.text.strip()
                if celltext.startswith('[') and celltext.endswith(']'):
                    text = celltext.strip('[]')
                    try:
                        sheet_code, cell_code = text.split('$')
                    except:
                        print('【错误】无法识别的内容', cell.text)
                        continue

                    # find sheet
                    sheet_name = None
                    for name in wb.sheetnames:
                        if sheet_code == name.split(' ')[0]:
                            sheet_name = name
                            break
                    if sheet_name is None:
                        print('【错误】无法找到工作表', sheet_code)
                        continue

                    # replace cell text
                    # cell.text = format_cell(wb[sheet_name][cell_code].value)
                    try:
                        replace_key_in_doc(cell, cell.text, format_cell(
                            wb[sheet_name][cell_code].value))
                    except Exception as e:
                        print('【错误】在替换', cell.text, '出错', e)
                        continue
    print()


def shuttle_text(shuttle):
    t = ''
    for i in shuttle:
        t += i.text
    return t


def replace_key_in_doc(doc, key, value):
    for p in doc.paragraphs:

        begin = 0
        for end in range(len(p.runs)):

            shuttle = p.runs[begin:end+1]

            full_text = shuttle_text(shuttle)
            # print('full_text:', full_text)
            if key in full_text:
                # print('Replace：', key, '->', value)
                # print([i.text for i in shuttle])

                # find the begin
                index = full_text.index(key)
                # print('full_text length', len(full_text), 'index:', index)
                while index >= len(p.runs[begin].text):
                    index -= len(p.runs[begin].text)
                    begin += 1

                shuttle = p.runs[begin:end+1]

                # do replace
                # print('before replace', [i.text for i in shuttle])
                if key in shuttle[0].text:
                    shuttle[0].text = shuttle[0].text.replace(key, value)
                else:
                    replace_begin_index = shuttle_text(shuttle).index(key)
                    replace_end_index = replace_begin_index + len(key)
                    replace_end_index_in_last_run = replace_end_index - \
                        len(shuttle_text(shuttle[:-1]))
                    shuttle[0].text = shuttle[0].text[:replace_begin_index] + value

                    # clear middle runs
                    for i in shuttle[1:-1]:
                        i.text = ''

                    # keep last run
                    shuttle[-1].text = shuttle[-1].text[replace_end_index_in_last_run:]

                # print('after replace', [i.text for i in shuttle])

                # set begin to next
                begin = end


if __name__ == '__main__':
    main()
